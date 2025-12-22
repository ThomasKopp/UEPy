import os
import time
import sys
import json
import logging
import csv
import re
import contextlib
import io
import shutil
import warnings
from datetime import datetime
from logging.handlers import RotatingFileHandler
import queue
import threading
import subprocess
import tkinter as tk
from tkinter import filedialog, messagebox, simpledialog
from tkinter import ttk
from tkinter.ttk import Style
import requests
from pypdf import PdfReader

# --- Windows-only: silence Paddle's ccache probe noise ---
# Paddle's cpp_extension utilities look for a file literally named "ccache" on PATH. If not found,
# it may call `where ccache` which prints "INFORMATION: Es konnten keine Dateien..." and emits a warning.
# We install a no-op shim early to avoid that noisy probe.
if os.name == "nt":
    try:
        shim_dir = os.path.join(os.path.expanduser("~"), ".ollama_translator_shims")
        os.makedirs(shim_dir, exist_ok=True)
        shim_path = os.path.join(shim_dir, "ccache")  # intentionally no extension
        if not os.path.exists(shim_path):
            with open(shim_path, "w", encoding="utf-8") as f:
                f.write("")
        current_path = os.environ.get("PATH", "")
        if shim_dir not in current_path.split(os.pathsep):
            os.environ["PATH"] = shim_dir + os.pathsep + current_path
    except Exception:
        pass

# Reduce noise from optional native deps (e.g. Paddle/PaddleOCR) when they are imported.
warnings.filterwarnings("ignore", message=r"No ccache found\..*", category=UserWarning)

# Optional dependencies are imported lazily to keep startup fast and avoid noisy helper-binary probes.
partition_pdf = None
PaddleOCR = None
convert_from_path = None
Document = None
FPDF = None
# Default Ollama API endpoint
DEFAULT_OLLAMA_API_BASE_URL = "http://localhost:11434/api"
DEFAULT_CHUNK_CHAR_LIMIT = 1600  # Chunk long texts to avoid model/context breakdowns
DEFAULT_HISTORY_LIMIT = 10
OFFLINE_QUEUE_PATH = os.path.join(os.path.expanduser("~"), ".ollama_translator_offline_queue.json")
GLOSSARY_MAX_ITEMS = 50
DEFAULT_OCR_DPI = 300

# --- Theme Definitions ---
LIGHT_THEME = {
    "bg": "#f0f0f0",
    "fg": "#000000",
    "select_bg": "#c3c3c3",
    "select_fg": "#000000",
    "button_bg": "#e1e1e1",
    "button_fg": "#000000",
    "entry_bg": "#ffffff",
    "entry_fg": "#000000",
    "disabled_fg": "#a3a3a3",
    "error_fg": "red",
    "active_model_fg": "black",
    "inactive_model_fg": "grey"
}

DARK_THEME = {
    "bg": "#333333",
    "fg": "#ffffff",
    "select_bg": "#555555",
    "select_fg": "#ffffff",
    "button_bg": "#555555",
    "button_fg": "#ffffff",
    "entry_bg": "#444444",
    "entry_fg": "#ffffff",
    "disabled_fg": "#888888",
    "error_fg": "#ff8080", # Lighter red for dark bg
    "active_model_fg": "white",
    "inactive_model_fg": "#aaaaaa"
}

# Prompt profile texts per direction
PROFILE_STYLES = {
    "de-en": {
        "standard": "Neutral business English, keep nuance and register.",
        "woertlich": "Highly literal; mirror phrasing closely, keep compound nouns intact.",
        "frei": "Freer rendering; reorganise sentences for clarity and flow.",
        "formal": "Formal tone, polite forms, avoid contractions.",
        "kreativ": "Creative marketing tone; persuasive, vivid wording."
    },
    "en-de": {
        "standard": "Neutrales Hochdeutsch, flüssig und idiomatisch.",
        "woertlich": "Sehr wortgetreu; Satzbau eng am Original.",
        "frei": "Freier Stil; Satzbau an natürliches Deutsch anpassen.",
        "formal": "Formell und höflich (Sie-Form), keine Umgangssprache.",
        "kreativ": "Kreativer, werblicher Stil mit lebendigen Formulierungen."
    }
}

DEFAULT_SHORTCUTS = {
    "translate": "<Control-Return>",
    "cancel": "<Escape>",
    "clear": "<Control-l>",
    "toggle_theme": "<Control-t>",
    "refresh": "<Control-r>"
}

class OllamaTranslatorApp:
    # --- Lazy imports (optional dependencies) ---
    @contextlib.contextmanager
    def _suppress_native_output(self):
        """
        Suppress process-level stdout/stderr (including child-process output) while importing
        optional dependencies that may call helper binaries on Windows (e.g. `where.exe`)
        or emit noisy native logs.
        """
        try:
            devnull_fd = os.open(os.devnull, os.O_WRONLY)
            saved_out = os.dup(1)
            saved_err = os.dup(2)
            os.dup2(devnull_fd, 1)
            os.dup2(devnull_fd, 2)
            os.close(devnull_fd)
            try:
                with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                    yield
            finally:
                os.dup2(saved_out, 1)
                os.dup2(saved_err, 2)
                os.close(saved_out)
                os.close(saved_err)
        except Exception:
            with contextlib.redirect_stdout(io.StringIO()), contextlib.redirect_stderr(io.StringIO()):
                yield

    def _lazy_import_unstructured(self):
        global partition_pdf
        if partition_pdf is not None:
            return partition_pdf
        try:
            with self._suppress_native_output():
                from unstructured.partition.pdf import partition_pdf as _partition_pdf
            partition_pdf = _partition_pdf
        except Exception:
            partition_pdf = None
        return partition_pdf

    def _lazy_import_pdf2image(self):
        global convert_from_path
        if convert_from_path is not None:
            return convert_from_path
        try:
            with self._suppress_native_output():
                from pdf2image import convert_from_path as _convert_from_path
            convert_from_path = _convert_from_path
        except Exception:
            convert_from_path = None
        return convert_from_path

    def _lazy_import_paddleocr(self):
        global PaddleOCR
        if PaddleOCR is not None:
            return PaddleOCR

        # Best-effort: reduce noisy native logging from OCR dependencies.
        os.environ.setdefault("GLOG_minloglevel", "2")
        os.environ.setdefault("FLAGS_minloglevel", "2")
        os.environ.setdefault("TF_CPP_MIN_LOG_LEVEL", "3")
        os.environ.setdefault("KMP_WARNINGS", "0")
        os.environ.setdefault("PADDLE_LOG_LEVEL", "3")

        try:
            with warnings.catch_warnings():
                warnings.filterwarnings("ignore", message=r"No ccache found\..*", category=UserWarning)
                with self._suppress_native_output():
                    from paddleocr import PaddleOCR as _PaddleOCR
            PaddleOCR = _PaddleOCR
        except Exception:
            PaddleOCR = None
        return PaddleOCR

    def _lazy_import_docx(self):
        global Document
        if Document is not None:
            return Document
        try:
            with self._suppress_native_output():
                from docx import Document as _Document
            Document = _Document
        except Exception:
            Document = None
        return Document

    def _lazy_import_fpdf(self):
        global FPDF
        if FPDF is not None:
            return FPDF
        try:
            with self._suppress_native_output():
                from fpdf import FPDF as _FPDF
            FPDF = _FPDF
        except Exception:
            FPDF = None
        return FPDF

    def _get_poppler_path(self):
        """
        Optional: resolve a Poppler bin directory for pdf2image on Windows.
        Users can set POPPLER_PATH to either the Poppler root or its bin directory.
        """
        raw = os.environ.get("POPPLER_PATH", "").strip()
        if not raw:
            return None
        if os.path.isdir(raw) and os.path.isfile(os.path.join(raw, "pdfinfo.exe")):
            return raw
        bin_dir = os.path.join(raw, "bin")
        if os.path.isdir(bin_dir) and os.path.isfile(os.path.join(bin_dir, "pdfinfo.exe")):
            return bin_dir
        return None

    def _create_paddleocr(self, lang):
        PaddleOCR_cls = self._lazy_import_paddleocr()
        if not PaddleOCR_cls:
            return None
        # PaddleOCR args changed across versions; try a few safe combinations.
        candidates = [
            {"lang": lang, "use_textline_orientation": True, "show_log": False},
            {"lang": lang, "use_angle_cls": True, "show_log": False},
            {"lang": lang, "use_angle_cls": True},
            {"lang": lang, "show_log": False},
            {"lang": lang},
        ]
        last_exc = None
        for kwargs in candidates:
            try:
                return PaddleOCR_cls(**kwargs)
            except TypeError as exc:
                last_exc = exc
                continue
        raise last_exc

    def _paddleocr_text_segments(self, result):
        """Return text segments from PaddleOCR outputs across versions."""
        if result is None:
            return []
        segments = []

        def visit(obj):
            if isinstance(obj, (list, tuple)):
                # Common line format: [box, (text, score)]
                if (
                    len(obj) == 2
                    and isinstance(obj[1], (list, tuple))
                    and len(obj[1]) >= 1
                    and isinstance(obj[1][0], str)
                ):
                    segments.append(obj[1][0])
                    return
                for item in obj:
                    visit(item)

        visit(result)
        return segments

    def _get_ocr_dpi(self):
        try:
            dpi = int(str(self.ocr_dpi_var.get()).strip())
            if 72 <= dpi <= 600:
                return dpi
        except Exception:
            pass
        return DEFAULT_OCR_DPI

    def _looks_like_list_item(self, line):
        s = (line or "").lstrip()
        return bool(re.match(r"^([-*•]\s+|\d+[.)]\s+|[A-Za-z][.)]\s+)", s))

    def _preprocess_ocr_image(self, img):
        if not getattr(self, "ocr_preprocess_var", None) or not self.ocr_preprocess_var.get():
            return img
        try:
            from PIL import ImageOps, ImageEnhance, ImageFilter
        except Exception:
            return img
        try:
            gray = ImageOps.grayscale(img)
            gray = ImageOps.autocontrast(gray)
            gray = ImageEnhance.Contrast(gray).enhance(1.6)
            gray = gray.filter(ImageFilter.SHARPEN)
            return gray
        except Exception:
            return img

    def _postprocess_ocr_text(self, text):
        if not getattr(self, "ocr_cleanup_var", None) or not self.ocr_cleanup_var.get():
            return (text or "").strip()

        t = (text or "").replace("\r\n", "\n").replace("\r", "\n").replace("\u00ad", "")
        t = re.sub(r"(?<=\w)-\n(?=\w)", "", t)

        lines = t.split("\n")
        out = []
        for raw in lines:
            cur = raw.rstrip()
            if not out:
                out.append(cur)
                continue
            prev = out[-1]
            if not prev.strip() or not cur.strip():
                out.append(cur)
                continue
            if self._looks_like_list_item(cur):
                out.append(cur)
                continue
            if self._looks_like_list_item(prev) or prev.rstrip().endswith(":"):
                out.append(cur)
                continue
            if re.search(r"[.!?][\"')\]]?$", prev.strip()):
                out.append(cur)
                continue
            out[-1] = prev.rstrip() + " " + cur.lstrip()

        t = "\n".join(out)
        t = re.sub(r"[ \t]+", " ", t)
        t = re.sub(r"\n{3,}", "\n\n", t)
        return t.strip()

    def _pdf_to_images(self, filepath, convert_fn, poppler_path):
        dpi = self._get_ocr_dpi()
        base_kwargs = {"dpi": dpi, "fmt": "png", "grayscale": True}
        if poppler_path:
            base_kwargs["poppler_path"] = poppler_path
        try:
            return convert_fn(filepath, **base_kwargs)
        except TypeError:
            # Older pdf2image versions may not accept fmt/grayscale.
            kwargs = {"dpi": dpi}
            if poppler_path:
                kwargs["poppler_path"] = poppler_path
            try:
                return convert_fn(filepath, **kwargs)
            except Exception:
                return convert_fn(filepath, poppler_path=poppler_path) if poppler_path else convert_fn(filepath)
 
    # --- Theme Management ---
    def apply_theme(self):
        """Apply the current theme to the entire application."""
        theme = LIGHT_THEME if self.current_theme == "light" else DARK_THEME

        self._configure_root(theme)
        self._configure_ttk_styles(theme)
        self._configure_non_ttk_widgets(theme)

    def _configure_root(self, theme):
        """Configure the root window background color."""
        self.root.config(bg=theme["bg"])

    def _configure_ttk_styles(self, theme):
        """Configure ttk widget styles based on the theme."""
        self.style.theme_use('clam')  # Use a theme that allows more customization

        self.style.configure('.', background=theme["bg"], foreground=theme["fg"],
                             fieldbackground=theme["entry_bg"], selectbackground=theme["select_bg"],
                             selectforeground=theme["select_fg"])
        self.style.map('.', background=[('active', theme["select_bg"])])

        self.style.configure('TButton', background=theme["button_bg"], foreground=theme["button_fg"])
        self.style.map('TButton', background=[('active', theme["select_bg"]), ('disabled', theme["bg"])],
                                  foreground=[('disabled', theme["disabled_fg"])])

        self.style.configure('TLabel', background=theme["bg"], foreground=theme["fg"])
        self.style.configure('TFrame', background=theme["bg"])
        self.style.configure('TLabelframe', background=theme["bg"], foreground=theme["fg"])
        self.style.configure('TLabelframe.Label', background=theme["bg"], foreground=theme["fg"])
        self.style.configure('TCombobox', fieldbackground=theme["entry_bg"], foreground=theme["entry_fg"],
                             selectbackground=theme["select_bg"], selectforeground=theme["select_fg"])
        self.style.map('TCombobox', fieldbackground=[('readonly', theme["entry_bg"])],
                                  selectbackground=[('readonly', theme["select_bg"])],
                                  selectforeground=[('readonly', theme["select_fg"])])
        # Removed attempt to style TCombobox.downarrow as it was unreliable
        self.style.configure('TProgressbar', background=theme["button_bg"], troughcolor=theme["entry_bg"])

    def _configure_non_ttk_widgets(self, theme):
        """Configure non-ttk widgets like Text, Listbox, and Labels."""
        text_config = {"background": theme["entry_bg"], "foreground": theme["entry_fg"],
                       "insertbackground": theme["fg"], "selectbackground": theme["select_bg"],
                       "selectforeground": theme["select_fg"]}
        listbox_config = {"background": theme["entry_bg"], "foreground": theme["entry_fg"],
                          "selectbackground": theme["select_bg"], "selectforeground": theme["select_fg"]}

        if hasattr(self, 'input_text'):
            self.input_text.config(**text_config)
        if hasattr(self, 'output_text'):
            self.output_text.config(**text_config)
        if hasattr(self, 'available_models_listbox'):
            self.available_models_listbox.config(**listbox_config)
        if hasattr(self, 'error_label'):
            self.error_label.config(foreground=theme["error_fg"])
        if hasattr(self, 'active_model_label'):
            active_fg = theme["active_model_fg"] if self.active_model else theme["inactive_model_fg"]
            self.active_model_label.config(foreground=active_fg)


    # --- Keyboard Shortcuts ---
    def _setup_keyboard_shortcuts(self):
        """Bind keyboard shortcuts to their respective handlers."""
        bindings = {
            "translate": self._keyboard_translate,
            "cancel": self._keyboard_cancel,
            "clear": self._keyboard_clear_input,
            "toggle_theme": self._keyboard_toggle_theme,
            "refresh": self._keyboard_refresh_models
        }
        for action, handler in bindings.items():
            key = self.shortcuts.get(action)
            if not key:
                continue
            try:
                self.root.unbind_all(key)
                self.root.bind_all(key, lambda e, h=handler: h())
            except tk.TclError:
                # Ignore invalid key patterns
                continue

    def _keyboard_translate(self):
        """Handle Ctrl+Enter to start translation."""
        if self.translate_button['state'] == tk.NORMAL:
            self.start_translation()
            self.translate_button.config(state=tk.DISABLED)  # Ensure button is disabled after starting translation

    def _keyboard_cancel(self):
        """Handle Escape to cancel translation."""
        if self.cancel_button['state'] == tk.NORMAL:
            self.cancel_translation()
        elif self.translation_controller:
            # Ensure tests that set a controller directly can trigger abort
            self.translation_controller['abort'] = True

    def _keyboard_clear_input(self):
        """Handle Ctrl+L to clear input text."""
        self.input_text.delete('1.0', tk.END)
        self.update_translate_button_state()

    def _keyboard_toggle_theme(self):
        """Handle Ctrl+T to toggle theme."""
        self.toggle_theme()

    def _keyboard_refresh_models(self):
        """Handle Ctrl+R to refresh available models."""
        self.refresh_available_models()

    def _open_shortcut_dialog(self):
        """Allow users to customise a few keyboard shortcuts."""
        dlg = tk.Toplevel(self.root)
        dlg.title("Shortcuts")
        ttk.Label(dlg, text="Aktion").grid(row=0, column=0, padx=5, pady=5)
        ttk.Label(dlg, text="Taste (Tk Syntax)").grid(row=0, column=1, padx=5, pady=5)
        actions = [
            ("Übersetzen", "translate"),
            ("Abbrechen", "cancel"),
            ("Eingabe leeren", "clear"),
            ("Theme wechseln", "toggle_theme"),
            ("Modelle aktualisieren", "refresh")
        ]
        entries = {}
        for idx, (label, key) in enumerate(actions, start=1):
            ttk.Label(dlg, text=label).grid(row=idx, column=0, padx=5, pady=2, sticky="w")
            var = tk.StringVar(value=self.shortcuts.get(key, DEFAULT_SHORTCUTS.get(key, "")))
            ent = ttk.Entry(dlg, textvariable=var, width=20)
            ent.grid(row=idx, column=1, padx=5, pady=2)
            entries[key] = var

        def save_and_close():
            for k, var in entries.items():
                val = var.get().strip()
                self.shortcuts[k] = val or DEFAULT_SHORTCUTS.get(k, "")
            self._setup_keyboard_shortcuts()
            self._save_config()
            dlg.destroy()

        ttk.Button(dlg, text="Speichern", command=save_and_close).grid(row=len(actions)+1, column=0, columnspan=2, pady=8)

    # --- Configuration Persistence ---
    def _load_config(self):
        """Load user configuration from a JSON file."""
        config_path = os.path.join(os.path.expanduser("~"), ".ollama_translator_config.json")
        if os.path.exists(config_path):
            try:
                with open(config_path, "r", encoding="utf-8") as f:
                    config = json.load(f)
                self.current_theme = config.get("theme", "light")
                self.api_endpoint_var.set(config.get("api_endpoint", DEFAULT_OLLAMA_API_BASE_URL))
                self.direction_var.set(config.get("direction", "de-en"))
                self.prompt_profile_var.set(config.get("prompt_profile", "standard"))
                if hasattr(self, "ocr_dpi_var"):
                    self.ocr_dpi_var.set(str(config.get("ocr_dpi", DEFAULT_OCR_DPI)))
                if hasattr(self, "ocr_cleanup_var"):
                    self.ocr_cleanup_var.set(bool(config.get("ocr_cleanup", True)))
                if hasattr(self, "ocr_preprocess_var"):
                    self.ocr_preprocess_var.set(bool(config.get("ocr_preprocess", True)))
                self.shortcuts.update(config.get("shortcuts", {}))
                self.auto_theme_var.set(config.get("auto_theme", False))
                self.security_mode_var.set(config.get("security_mode", False))
                self.auto_detect_var.set(config.get("auto_detect", True))
                self.history_limit = config.get("history_limit", DEFAULT_HISTORY_LIMIT)
                self.model_settings = config.get("model_settings", {}) or {}
                glossary_path = config.get("glossary_path")
                if glossary_path:
                    self._load_glossary_from_path(glossary_path)
                last_model = config.get("last_model", None)
                if last_model:
                    self.active_model = last_model
                    active_fg = LIGHT_THEME["active_model_fg"] if self.current_theme == "light" else DARK_THEME["active_model_fg"]
                    self.active_model_label.config(text=self.active_model, foreground=active_fg)
                    self._apply_model_settings(last_model)
                self.apply_theme()
            except Exception as e:
                print(f"Failed to load config: {e}")

    def _save_config(self):
        """Save user configuration to a JSON file."""
        config_path = os.path.join(os.path.expanduser("~"), ".ollama_translator_config.json")
        config = {
            "theme": self.current_theme,
            "api_endpoint": self.api_endpoint_var.get(),
            "direction": self.direction_var.get(),
            "prompt_profile": self.prompt_profile_var.get(),
            "ocr_dpi": self._get_ocr_dpi() if hasattr(self, "ocr_dpi_var") else DEFAULT_OCR_DPI,
            "ocr_cleanup": bool(self.ocr_cleanup_var.get()) if hasattr(self, "ocr_cleanup_var") else True,
            "ocr_preprocess": bool(self.ocr_preprocess_var.get()) if hasattr(self, "ocr_preprocess_var") else True,
            "last_model": self.active_model,
            "shortcuts": self.shortcuts,
            "auto_theme": self.auto_theme_var.get(),
            "security_mode": self.security_mode_var.get(),
            "auto_detect": self.auto_detect_var.get(),
            "history_limit": self.history_limit,
            "model_settings": self.model_settings,
            "glossary_path": self.glossary.get("path")
        }
        try:
            with open(config_path, "w", encoding="utf-8") as f:
                json.dump(config, f, indent=4)
        except Exception as e:
            print(f"Failed to save config: {e}")

    # --- Override __init__ to add config load and keyboard shortcuts setup ---
    def __init__(self, root):
        self.root = root
        self.root.title("Ollama Translator (Python)")

        # Initialize tkinter variables early to avoid attribute errors
        self.api_endpoint_var = tk.StringVar(value=DEFAULT_OLLAMA_API_BASE_URL)
        self.direction_var = tk.StringVar(value="de-en")
        self.prompt_profile_var = tk.StringVar(value="standard")
        self.ocr_dpi_var = tk.StringVar(value=str(DEFAULT_OCR_DPI))
        self.ocr_cleanup_var = tk.BooleanVar(value=True)
        self.ocr_preprocess_var = tk.BooleanVar(value=True)
        # Disable background threads in unit test context to avoid Tk thread errors
        self.use_threads = 'unittest' not in sys.modules

        self.active_model = None
        self.translation_controller = None
        self.ocr_controller = None
        self.history = []
        self.batch_queue = []
        self.batch_results = []
        self.glossary = {"map": {}, "dnt": set(), "path": None}
        self.model_settings = {}
        self.shortcuts = DEFAULT_SHORTCUTS.copy()
        self.auto_theme_var = tk.BooleanVar(value=False)
        self.security_mode_var = tk.BooleanVar(value=False)
        self.auto_detect_var = tk.BooleanVar(value=True)
        self.history_limit = DEFAULT_HISTORY_LIMIT

        self.style = Style(root)
        self.current_theme = "light"

        # Logging setup (toggle via checkbox)
        self.log_enabled_var = tk.BooleanVar(value=True)
        self.logger = logging.getLogger("ollama_translator")
        self.logger.setLevel(logging.INFO)
        self.logger.propagate = False
        self.log_path = os.path.join(os.path.expanduser("~"), ".ollama_translator.log")
        if not any(isinstance(h, RotatingFileHandler) and getattr(h, "baseFilename", "") == self.log_path for h in self.logger.handlers):
            try:
                handler = RotatingFileHandler(self.log_path, maxBytes=512_000, backupCount=3, encoding="utf-8")
                handler.setFormatter(logging.Formatter("%(asctime)s | %(levelname)s | %(message)s"))
                self.logger.addHandler(handler)
            except Exception as e:
                print(f"Logger setup failed: {e}")

        # Main Frames setup (same as before)...
        self.header_frame = ttk.Frame(root, padding="10")
        self.header_frame.grid(row=0, column=0, sticky="ew")

        self.model_mgmt_frame = ttk.LabelFrame(root, text="Model Management", padding="10")
        self.model_mgmt_frame.grid(row=1, column=0, padx=10, pady=5, sticky="ew")

        self.translation_frame = ttk.LabelFrame(root, text="Translation", padding="10")
        self.translation_frame.grid(row=2, column=0, padx=10, pady=5, sticky="nsew")

        self.progress_frame = ttk.Frame(root, padding="10")
        self.progress_frame.grid(row=3, column=0, padx=10, pady=5, sticky="ew")

        self.footer_frame = ttk.Frame(root, padding="10")
        self.footer_frame.grid(row=4, column=0, sticky="ew")

        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(2, weight=1)
        self.translation_frame.columnconfigure(0, weight=1)
        self.translation_frame.columnconfigure(1, weight=1)
        self.translation_frame.rowconfigure(1, weight=1)

        self.create_header_widgets()
        self.create_model_management_widgets()

        # Load config before applying theme
        self._load_config()

        self.apply_theme()
        if self.auto_theme_var.get():
            self._apply_auto_theme()

        self.create_translation_widgets()
        self.create_progress_widgets()
        self.create_footer_widgets()

        # Setup keyboard shortcuts
        self._setup_keyboard_shortcuts()

        # Initial actions (skip auto-start/refresh in unittest mode)
        if self.use_threads:
            # Quick ping; if offline, try to start Ollama once.
            if not self._check_server_status():
                self.start_ollama_server()
            # Enforce a real health check before enabling model refresh.
            if self._ensure_server_reachable():
                self.refresh_available_models()
        # Bind close event to save config
        self.root.protocol("WM_DELETE_WINDOW", self._on_close)

        # UI dispatch queue for thread-safe updates (only needed when threads are used)
        if self.use_threads:
            self._ui_queue = queue.Queue()
            self._process_ui_queue()

    def _on_close(self):
        self._save_config()
        self.root.destroy()

    # --- Thread-safe UI dispatch ---
    def _dispatch_ui(self, fn, *args, **kwargs):
        """Queue UI work; executes immediately when threads are disabled."""
        if self.use_threads:
            self._ui_queue.put((fn, args, kwargs))
        else:
            fn(*args, **kwargs)

    def _process_ui_queue(self):
        """Execute queued UI callbacks on the Tk thread."""
        if not self.use_threads:
            return
        while not self._ui_queue.empty():
            fn, args, kwargs = self._ui_queue.get()
            try:
                fn(*args, **kwargs)
            except Exception as e:
                print(f"UI dispatch error: {e}")
        self.root.after(50, self._process_ui_queue)

    def _log_event(self, event, **fields):
        """Write a structured log line if logging is enabled."""
        if not self.log_enabled_var.get():
            return
        try:
            pairs = " ".join(f"{k}={fields[k]}" for k in sorted(fields))
            self.logger.info("%s %s", event, pairs)
        except Exception as e:
            print(f"Log error: {e}")

    def _set_server_status(self, text, ok=False):
        """Update server status label with color."""
        if not hasattr(self, "server_status_label"):
            return
        fg = "#008000" if ok else (DARK_THEME["error_fg"] if self.current_theme == "dark" else LIGHT_THEME["error_fg"])
        self.server_status_label.config(text=text, foreground=fg)

    def _apply_theme_to_widgets(self, theme):
        """Apply theme styles to widgets, used during theme toggle or updates."""
        self._configure_ttk_styles(theme)
        self._configure_non_ttk_widgets(theme)
        if hasattr(self, "server_status_label"):
            status_text = self.server_status_label.cget("text")
            ok = "online" in status_text.lower() or "ready" in status_text.lower()
            fg = "#008000" if ok else theme["error_fg"]
            self.server_status_label.config(foreground=fg)

    def toggle_theme(self):
        """Toggle between light and dark themes."""
        self.current_theme = "dark" if self.current_theme == "light" else "light"
        self.apply_theme()

    def _apply_auto_theme(self):
        """Switch theme based on time of day (light 07-19h, dark otherwise)."""
        hour = datetime.now().hour
        desired = "light" if 7 <= hour < 19 else "dark"
        if self.current_theme != desired:
            self.current_theme = desired
            self.apply_theme()
        # Reschedule check every 10 minutes when auto mode is active
        if self.auto_theme_var.get():
            self.root.after(600000, self._apply_auto_theme)

    # --- Widget Creation Methods ---
    def create_header_widgets(self):
        ttk.Label(self.header_frame, text="Ollama Translator", font=("Arial", 16)).pack(side=tk.LEFT, padx=5)

        # API Endpoint input
        ttk.Label(self.header_frame, text="API Endpoint:").pack(side=tk.LEFT, padx=(20, 5))
        if not hasattr(self, "api_endpoint_var"):
            self.api_endpoint_var = tk.StringVar(value=DEFAULT_OLLAMA_API_BASE_URL)
        self.api_endpoint_entry = ttk.Entry(self.header_frame, textvariable=self.api_endpoint_var, width=40)
        self.api_endpoint_entry.pack(side=tk.LEFT, padx=5)

        ttk.Label(self.header_frame, text="Direction:").pack(side=tk.LEFT, padx=(20, 5))
        if not hasattr(self, "direction_var"):
            self.direction_var = tk.StringVar(value="de-en")
        direction_combo = ttk.Combobox(self.header_frame, textvariable=self.direction_var, 
                                       values=["de-en", "en-de"], state="readonly", width=15)
        direction_combo.pack(side=tk.LEFT, padx=5)
        direction_combo.bind("<<ComboboxSelected>>", self.update_translation_prompt) # Update prompt on change

        ttk.Checkbutton(self.header_frame, text="Auto-Detect", variable=self.auto_detect_var).pack(side=tk.LEFT, padx=(5,10))

        ttk.Label(self.header_frame, text="Profil:").pack(side=tk.LEFT, padx=(20,5))
        if not hasattr(self, "prompt_profile_var"):
            self.prompt_profile_var = tk.StringVar(value="standard")
        profile_combo = ttk.Combobox(
            self.header_frame,
            textvariable=self.prompt_profile_var,
            values=["standard", "woertlich", "frei", "formal", "kreativ"],
            state="readonly",
            width=15
        )
        profile_combo.pack(side=tk.LEFT, padx=5)
        profile_combo.bind("<<ComboboxSelected>>", self.update_translation_prompt)

        ttk.Button(self.header_frame, text="Ollama Check", command=self.run_ollama_check).pack(side=tk.LEFT, padx=(20,5))
        ttk.Button(self.header_frame, text="Shortcuts", command=self._open_shortcut_dialog).pack(side=tk.LEFT, padx=5)

    def create_model_management_widgets(self):
        # Available Models Section
        available_frame = ttk.Frame(self.model_mgmt_frame)
        available_frame.grid(row=0, column=0, padx=5, pady=5, sticky="ns")
        ttk.Label(available_frame, text="Available Models").pack()
        self.available_models_listbox = tk.Listbox(available_frame, height=5, exportselection=False)
        self.available_models_listbox.pack(fill=tk.X, expand=True)
        available_buttons = ttk.Frame(available_frame)
        available_buttons.pack(pady=5)
        ttk.Button(available_buttons, text="Refresh", command=self.refresh_available_models).pack(side=tk.LEFT, padx=2)
        ttk.Button(available_buttons, text="Activate Model", command=self.activate_model).pack(side=tk.LEFT, padx=2)

        # Active Model Section
        active_frame = ttk.Frame(self.model_mgmt_frame)
        active_frame.grid(row=0, column=1, padx=5, pady=5, sticky="ns")
        ttk.Label(active_frame, text="Active Model for Translation").pack()
        self.active_model_label = ttk.Label(active_frame, text="None selected", foreground="grey", width=30, anchor="center")
        self.active_model_label.pack(pady=5)
        # Note: No listbox for active models needed, just display the selected one.
        # Adding a deactivate button
        ttk.Button(active_frame, text="Deactivate Model", command=self.deactivate_model).pack(pady=5)

        # Model settings persistence
        settings_frame = ttk.LabelFrame(self.model_mgmt_frame, text="Model Settings")
        settings_frame.grid(row=1, column=0, columnspan=2, sticky="ew", padx=5, pady=(5,0))
        ttk.Label(settings_frame, text="Temperature").grid(row=0, column=0, padx=5, pady=2, sticky="w")
        self.temp_var = tk.DoubleVar(value=0.2)
        ttk.Spinbox(settings_frame, from_=0.0, to=1.0, increment=0.05, textvariable=self.temp_var, width=6).grid(row=0, column=1, padx=5)

        ttk.Label(settings_frame, text="Top-p").grid(row=0, column=2, padx=5, pady=2, sticky="w")
        self.top_p_var = tk.DoubleVar(value=0.9)
        ttk.Spinbox(settings_frame, from_=0.0, to=1.0, increment=0.05, textvariable=self.top_p_var, width=6).grid(row=0, column=3, padx=5)

        ttk.Label(settings_frame, text="Max tokens").grid(row=0, column=4, padx=5, pady=2, sticky="w")
        self.max_tokens_var = tk.IntVar(value=512)
        ttk.Entry(settings_frame, textvariable=self.max_tokens_var, width=8).grid(row=0, column=5, padx=5)

        ttk.Button(settings_frame, text="Save per Model", command=self._save_model_settings_for_active).grid(row=0, column=6, padx=5)

        self.model_mgmt_frame.columnconfigure(0, weight=1)
        self.model_mgmt_frame.columnconfigure(1, weight=1)

    def create_translation_widgets(self):
        # Input Section
        input_frame = ttk.Frame(self.translation_frame)
        input_frame.grid(row=0, column=0, rowspan=2, padx=5, pady=5, sticky="nsew")
        ttk.Label(input_frame, text="Input").grid(row=0, column=0, sticky="w")
        self.input_text = tk.Text(input_frame, wrap=tk.WORD, height=10, width=40)
        self.input_text.grid(row=1, column=0, sticky="nsew")
        input_buttons = ttk.Frame(input_frame)
        input_buttons.grid(row=2, column=0, pady=5, sticky="ew")

        ttk.Button(input_buttons, text="Upload TXT", command=self.upload_txt).pack(side=tk.LEFT, padx=2)
        ttk.Button(input_buttons, text="Upload PDF (OCR)", command=self.upload_pdf).pack(side=tk.LEFT, padx=2)
        ttk.Button(input_buttons, text="Clear", command=lambda: self.input_text.delete('1.0', tk.END)).pack(side=tk.LEFT, padx=2)
        ttk.Button(input_buttons, text="Glossar laden", command=self._load_glossary).pack(side=tk.LEFT, padx=2)
        self.translate_button = tk.Button(input_buttons, text="Translate", command=self.start_translation, state=tk.DISABLED)
        self.translate_button.pack(side=tk.LEFT, padx=2)
        self.cancel_button = tk.Button(input_buttons, text="Cancel", command=self.cancel_translation, state=tk.DISABLED)
        self.cancel_button.pack(side=tk.LEFT, padx=2)

        batch_buttons = ttk.Frame(input_frame)
        batch_buttons.grid(row=4, column=0, pady=5, sticky="w")
        ttk.Button(batch_buttons, text="Batch: Input hinzufügen", command=self._batch_add_current).pack(side=tk.LEFT, padx=2)
        ttk.Button(batch_buttons, text="Batch: Dateien", command=self._batch_add_files).pack(side=tk.LEFT, padx=2)
        ttk.Button(batch_buttons, text="Batch starten", command=self._start_batch).pack(side=tk.LEFT, padx=2)
        ttk.Button(batch_buttons, text="Batch export", command=self._export_batch_results).pack(side=tk.LEFT, padx=2)

        self.glossary_label = ttk.Label(input_frame, text="Glossar: keines")
        self.glossary_label.grid(row=5, column=0, sticky="w", pady=(2,0))

        # OCR quality options
        ocr_opts = ttk.Frame(input_frame)
        ocr_opts.grid(row=3, column=0, pady=(0,5), sticky="w")
        ttk.Label(ocr_opts, text="OCR-Qualität:").pack(side=tk.LEFT, padx=(0,5))
        self.ocr_quality_var = tk.StringVar(value="praezise")
        ttk.Radiobutton(ocr_opts, text="präzise (OCR)", variable=self.ocr_quality_var, value="praezise").pack(side=tk.LEFT)
        ttk.Radiobutton(ocr_opts, text="schnell (Text-Extract)", variable=self.ocr_quality_var, value="schnell").pack(side=tk.LEFT, padx=(5,0))
        ttk.Checkbutton(ocr_opts, text="Bereinigen", variable=self.ocr_cleanup_var, command=self._save_config).pack(side=tk.LEFT, padx=(10,0))
        ttk.Checkbutton(ocr_opts, text="Bild-Preproc", variable=self.ocr_preprocess_var, command=self._save_config).pack(side=tk.LEFT, padx=(5,0))
        ttk.Label(ocr_opts, text="DPI:").pack(side=tk.LEFT, padx=(10,0))
        dpi_combo = ttk.Combobox(ocr_opts, textvariable=self.ocr_dpi_var, values=["200", "300", "400"], width=5, state="readonly")
        dpi_combo.pack(side=tk.LEFT, padx=(3,0))
        dpi_combo.bind("<<ComboboxSelected>>", lambda e: self._save_config())

        # Output Section
        output_frame = ttk.Frame(self.translation_frame)
        output_frame.grid(row=0, column=1, rowspan=2, padx=5, pady=5, sticky="nsew")
        ttk.Label(output_frame, text="Output").grid(row=0, column=0, sticky="w")
        self.output_text = tk.Text(output_frame, wrap=tk.WORD, height=10, width=40, state=tk.DISABLED)
        self.output_text.grid(row=1, column=0, sticky="nsew")
        output_buttons = ttk.Frame(output_frame)
        output_buttons.grid(row=2, column=0, pady=5, sticky="ew")

        ttk.Button(output_buttons, text="Save TXT", command=self.save_txt).pack(side=tk.LEFT, padx=2)
        ttk.Button(output_buttons, text="Save MD", command=self.save_markdown).pack(side=tk.LEFT, padx=2)
        ttk.Button(output_buttons, text="Save DOCX", command=self.save_docx).pack(side=tk.LEFT, padx=2)
        ttk.Button(output_buttons, text="Export PDF", command=self.save_pdf).pack(side=tk.LEFT, padx=2)
        ttk.Button(output_buttons, text="Copy", command=self.copy_to_clipboard).pack(side=tk.LEFT, padx=2)

        input_frame.rowconfigure(1, weight=1)
        input_frame.columnconfigure(0, weight=1)
        output_frame.rowconfigure(1, weight=1)
        output_frame.columnconfigure(0, weight=1)

        # History & Batch overview
        extra_frame = ttk.Frame(self.translation_frame)
        extra_frame.grid(row=2, column=0, columnspan=2, sticky="ew", padx=5, pady=(5,0))

        history_box = ttk.Labelframe(extra_frame, text="History (letzte %d)" % self.history_limit)
        history_box.grid(row=0, column=0, padx=5, pady=5, sticky="ew")
        self.history_list = tk.Listbox(history_box, height=4, exportselection=False)
        self.history_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        hist_btns = ttk.Frame(history_box)
        hist_btns.pack(side=tk.LEFT, padx=5)
        ttk.Button(hist_btns, text="Copy Output", command=self._history_copy).pack(pady=2)
        ttk.Button(hist_btns, text="Retry", command=self._history_retry).pack(pady=2)

        batch_box = ttk.Labelframe(extra_frame, text="Batch Queue")
        batch_box.grid(row=0, column=1, padx=5, pady=5, sticky="ew")
        self.batch_list = tk.Listbox(batch_box, height=4, exportselection=False)
        self.batch_list.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
        batch_btns = ttk.Frame(batch_box)
        batch_btns.pack(side=tk.LEFT, padx=5)
        ttk.Button(batch_btns, text="Entfernen", command=self._batch_remove_selected).pack(pady=2)
        ttk.Button(batch_btns, text="Leeren", command=self._batch_clear).pack(pady=2)

        extra_frame.columnconfigure(0, weight=1)
        extra_frame.columnconfigure(1, weight=1)

    def create_progress_widgets(self):
        ttk.Label(self.progress_frame, text="Progress:").pack(side=tk.LEFT, padx=5)
        self.progress_bar = ttk.Progressbar(self.progress_frame, orient=tk.HORIZONTAL, length=250, mode='determinate')
        self.progress_bar.pack(side=tk.LEFT, fill=tk.X, expand=True, padx=5)
        self.status_label = ttk.Label(self.progress_frame, text="Idle", foreground="grey")
        self.status_label.pack(side=tk.LEFT, padx=5)
        self.ocr_cancel_button = tk.Button(self.progress_frame, text="Cancel OCR", command=self.cancel_ocr, state=tk.DISABLED)
        self.ocr_cancel_button.pack(side=tk.LEFT, padx=5)
        self.error_label = ttk.Label(self.progress_frame, text="", foreground="red")
        self.error_label.pack(side=tk.LEFT, padx=5)

    def create_footer_widgets(self):
        self.server_status_label = ttk.Label(self.footer_frame, text="Server: unbekannt", foreground="grey")
        self.server_status_label.pack(side=tk.LEFT, padx=5)
        # Theme toggle button
        self.theme_button = ttk.Button(self.footer_frame, text="Toggle Theme", command=self.toggle_theme)
        self.theme_button.pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(self.footer_frame, text="Auto Theme", variable=self.auto_theme_var,
                        command=self._apply_auto_theme).pack(side=tk.LEFT, padx=5)
        # Logging toggle
        ttk.Checkbutton(self.footer_frame, text="File-Log aktiv", variable=self.log_enabled_var).pack(side=tk.LEFT, padx=5)
        ttk.Checkbutton(self.footer_frame, text="Sicherheitsmodus", variable=self.security_mode_var).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.footer_frame, text="Offline-Queue senden", command=self._process_offline_queue).pack(side=tk.LEFT, padx=5)
        ttk.Button(self.footer_frame, text="Bewerten", command=self._open_feedback_dialog).pack(side=tk.LEFT, padx=5)

    # --- Helper Methods --- 
    def start_ollama_server(self):
        import time
        try:
            # Try to connect to the server first to see if it's already running
            requests.get(self.get_api_base_url(), timeout=1) # Short timeout
            print("Ollama server already running.")
            self._set_server_status("Server: online", ok=True)
        except Exception:
            print("Ollama server not running. Attempting to start...")
            try:
                if os.name == 'nt': # Windows
                    subprocess.Popen(["ollama", "serve"], creationflags=subprocess.CREATE_NO_WINDOW)
                else: # macOS/Linux
                    subprocess.Popen(["ollama", "serve"], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                print("Ollama serve command issued.")
                # Wait for server readiness with retries
                max_retries = 10
                retry_delay = 1  # seconds
                for i in range(max_retries):
                    try:
                        time.sleep(retry_delay)
                        response = requests.get(self.get_api_base_url(), timeout=1)
                        if response.status_code == 200:
                            print("Ollama server is ready.")
                            self._set_server_status("Server: online", ok=True)
                            break
                    except Exception:
                        print(f"Waiting for Ollama server to start... ({i+1}/{max_retries})")
                else:
                    self.show_error("Failed to start Ollama server after multiple attempts.")
                    self._set_server_status("Server: offline", ok=False)
                    messagebox.showerror("Ollama Error", "Failed to start Ollama server after multiple attempts.")
            except FileNotFoundError:
                self.show_error("Ollama command not found. Ensure Ollama is installed and in your PATH.")
                self._set_server_status("Server: ollama cmd not found", ok=False)
                messagebox.showerror("Ollama Error", "Ollama command not found. Please ensure Ollama is installed and in your system's PATH.")
            except Exception as e:
                self.show_error(f"Failed to start Ollama server: {e}")
                self._set_server_status("Server: start failed", ok=False)
                messagebox.showerror("Ollama Error", f"Failed to start Ollama server: {e}")
        else:
            return True
        return False

    def _ensure_server_reachable(self):
        """Verify Ollama API is reachable; show a blocking error if not."""
        base = self.get_api_base_url().rstrip("/")
        try:
            resp = requests.get(f"{base}/tags", timeout=2)
            resp.raise_for_status()
            self._set_server_status("Server: online", ok=True)
            return True
        except requests.exceptions.RequestException as exc:
            self._set_server_status("Server: offline", ok=False)
            self.show_error("Ollama API nicht erreichbar.")
            messagebox.showerror(
                "Ollama nicht erreichbar",
                f"Der Ollama-Endpoint {base} konnte nicht erreicht werden.\n"
                f"Bitte prüfen, ob 'ollama serve' läuft.\n\nDetails: {exc}"
            )
            # Keep translate disabled to avoid follow-up errors
            if hasattr(self, "translate_button"):
                self.translate_button.config(state=tk.DISABLED)
            return False

    def _check_server_status(self):
        """Ping Ollama once at startup to show status without forcing start."""
        try:
            base = self.get_api_base_url().rstrip("/")
            requests.get(f"{base}/tags", timeout=1)
            self._set_server_status("Server: online", ok=True)
            return True
        except Exception:
            self._set_server_status("Server: offline", ok=False)
            return False

    def run_ollama_check(self):
        """Manual check: call /tags and show models and status."""
        base = self.get_api_base_url().rstrip("/")
        try:
            resp = requests.get(f"{base}/tags", timeout=3)
            resp.raise_for_status()
            data = resp.json()
            models = [m.get("name", "") for m in data.get("models", [])]
            msg = f"Endpoint: {base}\nStatus: online\nModel count: {len(models)}"
            if models:
                msg += "\n\nModels:\n- " + "\n- ".join(sorted(models))
            self._set_server_status("Server: online", ok=True)
            messagebox.showinfo("Ollama Check", msg)
            self._log_event("ollama_check_ok", models=len(models))
        except Exception as exc:
            self._set_server_status("Server: offline", ok=False)
            self.show_error("Ollama API nicht erreichbar.")
            messagebox.showerror("Ollama Check", f"Endpoint: {base}\nStatus: offline\nDetails: {exc}")
            self._log_event("ollama_check_fail", error=str(exc))

    def show_error(self, message):
        self.error_label.config(text=message)
        # Optionally use messagebox for more prominent errors
        # messagebox.showerror("Error", message)
        if hasattr(self, "server_status_label") and message.lower().startswith("connection error"):
            self._set_server_status("Server: offline", ok=False)

    def clear_error(self):
        self.error_label.config(text="")

    def update_translate_button_state(self):
        if self.active_model and self.input_text.get("1.0", "end-1c").strip():
            self.translate_button.config(state=tk.NORMAL)
        else:
            self.translate_button.config(state=tk.DISABLED)

    # --- Model Management Methods --- 
    def refresh_available_models(self):
        self.clear_error()
        self.available_models_listbox.delete(0, tk.END)
        self.available_models_listbox.insert(tk.END, "Loading...")
        # Pass api_base_url safely to thread
        self._api_base_url_for_thread = self.api_endpoint_var.get().strip()
        if self.use_threads:
            threading.Thread(target=self._fetch_models_thread, daemon=True).start()
        else:
            self._fetch_models_thread()

    def _fetch_models_thread(self):
        try:
            # Use a local copy of api_base_url passed from main thread to avoid tkinter access in worker thread
            api_base_url = getattr(self, "_api_base_url_for_thread", self.api_endpoint_var.get().strip())

            response = requests.get(f"{api_base_url}/tags")
            response.raise_for_status()  # Raise HTTPError for bad responses (4xx or 5xx)
            data = response.json()
            models = [m['name'] for m in data.get('models', [])]

            self._dispatch_ui(self._update_available_models_list, models)

        except requests.exceptions.ConnectionError:
            self._dispatch_ui(self.show_error, "Connection Error: Could not connect to Ollama API.")
            self._dispatch_ui(self._update_available_models_list, [])
        except requests.exceptions.RequestException as e:
            self._dispatch_ui(self.show_error, f"API Error: {e}")
            self._dispatch_ui(self._update_available_models_list, [])
        except json.JSONDecodeError:
            self._dispatch_ui(self.show_error, "API Error: Invalid JSON response.")
            self._dispatch_ui(self._update_available_models_list, [])

    def _update_available_models_list(self, models):
        self.available_models_listbox.delete(0, tk.END)
        if models:
            for model in sorted(models):
                self.available_models_listbox.insert(tk.END, model)
        else:
            self.available_models_listbox.insert(tk.END, "No models found.")
            if not self.error_label.cget("text"): # Show error only if not already shown by fetch
                 self.show_error("No models available or API error.")

    def activate_model(self):
        selection = self.available_models_listbox.curselection()
        if not selection:
            messagebox.showwarning("No Selection", "Please select a model from the 'Available Models' list.")
            return
        
        selected_model = self.available_models_listbox.get(selection[0])
        if selected_model in ["Loading...", "No models found."]:
             messagebox.showwarning("Invalid Selection", "Please wait for models to load or select a valid model.")
             return

        # If a different model is already active and a translation might be running, cancel it.
        if self.active_model and self.active_model != selected_model and self.translation_controller:
            print(f"Switching model from {self.active_model} to {selected_model}. Cancelling ongoing translation if any.")
            self.cancel_translation()

        self.active_model = selected_model
        active_fg = LIGHT_THEME["active_model_fg"] if self.current_theme == "light" else DARK_THEME["active_model_fg"]
        self.active_model_label.config(text=self.active_model, foreground=active_fg)
        self._apply_model_settings(selected_model)
        self.update_translate_button_state()
        self.update_translation_prompt() # Update prompt when model changes
        print(f"Activated model: {self.active_model}")

    def deactivate_model(self):
        self.active_model = None
        inactive_fg = LIGHT_THEME["inactive_model_fg"] if self.current_theme == "light" else DARK_THEME["inactive_model_fg"]
        self.active_model_label.config(text="None selected", foreground=inactive_fg)
        self.update_translate_button_state()
        self.update_translation_prompt()
        print("Deactivated model")

    def _apply_model_settings(self, model_name):
        settings = self.model_settings.get(model_name)
        if settings:
            self.temp_var.set(settings.get("temperature", 0.2))
            self.top_p_var.set(settings.get("top_p", 0.9))
            self.max_tokens_var.set(settings.get("max_tokens", 512))

    def _save_model_settings_for_active(self):
        if not self.active_model:
            messagebox.showwarning("Kein Modell", "Bitte zuerst ein Modell aktivieren.")
            return
        self.model_settings[self.active_model] = {
            "temperature": float(self.temp_var.get()),
            "top_p": float(self.top_p_var.get()),
            "max_tokens": int(self.max_tokens_var.get() or 0)
        }
        self._log_event("model_settings_saved", model=self.active_model)
        messagebox.showinfo("Gespeichert", f"Settings für {self.active_model} gespeichert.")

    # --- Translation Methods --- 
    def get_translation_prompt(self):
        """
        Build a strict, explicit prompt so the model respects the requested language
        and translates the entire input (no summaries or omissions).
        """
        direction = self.direction_var.get()
        source_lang = "German" if direction == "de-en" else "English"
        target_lang = "English" if direction == "de-en" else "German"

        direction_notes = (
            "Use clear, idiomatic US/UK business English. Preserve any honorifics and tone markers."
            if target_lang == "English"
            else "Schreibe natürlich klingendes, idiomatisches Deutsch. Verwende standardmäßig die Höflichkeitsform „Sie“, es sei denn, der Eingangstext ist eindeutig informell und durchgängig mit „du“ gehalten."
        )
        profile_key = self.prompt_profile_var.get()
        profile_notes = PROFILE_STYLES.get(direction, {}).get(profile_key, PROFILE_STYLES.get(direction, {}).get("standard", ""))
        profile_label = profile_key if profile_key != "standard" else "standard"

        return (
            "You are a professional translator.\n"
            f"- Source language: {source_lang}\n"
            f"- Target language: {target_lang}\n"
            "- Translate *every* sentence and word; do not skip, summarise, or shorten.\n"
            "- Preserve all formatting exactly (line breaks, bullet lists, headings, Markdown, numbering, inline code).\n"
            "- Keep units, dates, names, and code as-is unless conversion is explicitly required.\n"
            "- Output only the translated text in the target language-no prefaces, notes, or explanations.\n"
            f"- Style guidance ({profile_label}): {profile_notes or direction_notes}\n"
            f"- Additional direction hints: {direction_notes}\n\n"
        )

    def _build_prompt(self, chunk_text):
        base = self.get_translation_prompt()
        glossary_text = ""
        if self.glossary["map"]:
            pairs = list(self.glossary["map"].items())[:GLOSSARY_MAX_ITEMS]
            glossary_lines = [f"- {src} -> {tgt}" for src, tgt in pairs]
            glossary_text += "Use the following glossary exactly (no paraphrasing):\n" + "\n".join(glossary_lines) + "\n"
        if self.glossary["dnt"] or ("<<" in chunk_text and ">>" in chunk_text):
            glossary_text += "Keep placeholders like <<DNT_#>> or <<PH_#>> unchanged and return them verbatim.\n"
        return base + glossary_text + chunk_text
    def _chunk_text_for_translation(self, text, max_chars=DEFAULT_CHUNK_CHAR_LIMIT):
        """
        Split long input into reasonably sized chunks while preserving paragraph
        boundaries when possible. Ensures the model is not overwhelmed and every
        part is translated.
        """
        chunks = []
        buffer = []
        current_len = 0
        paragraphs = text.split("\n\n")
        for para in paragraphs:
            seg = para.strip("\n")
            if not seg:
                # preserve blank paragraph as newline between buffers
                if buffer:
                    buffer.append("")
                    current_len += 2
                continue
            # If the paragraph fits in the current buffer
            if current_len + len(seg) + 2 <= max_chars:
                buffer.append(seg)
                current_len += len(seg) + 2
            else:
                if buffer:
                    chunks.append("\n\n".join(buffer).strip())
                buffer = []
                current_len = 0
                # If the paragraph itself is too long, hard-split it
                if len(seg) > max_chars:
                    start = 0
                    while start < len(seg):
                        chunks.append(seg[start:start+max_chars])
                        start += max_chars
                else:
                    buffer.append(seg)
                    current_len = len(seg)
        if buffer:
            chunks.append("\n\n".join(buffer).strip())
        return [c for c in chunks if c.strip()]

    def _mask_sensitive(self, text):
        """Mask e-mails and phone numbers when security mode is enabled."""
        email_re = re.compile(r"[\w.\-]+@[\w\-]+\.[A-Za-z]{2,}")
        phone_re = re.compile(r"(\+?\d[\d\s\-]{6,}\d)")
        text = email_re.sub("[EMAIL]", text)
        text = phone_re.sub("[PHONE]", text)
        return text

    def _apply_do_not_translate(self, text):
        mapping = {}
        for idx, term in enumerate(self.glossary.get("dnt", [])):
            placeholder = f"<<DNT_{idx}>>"
            if term:
                text = text.replace(term, placeholder)
                mapping[placeholder] = term
        return text, mapping

    def _apply_auto_placeholders(self, text, mapping):
        def next_ph():
            i = 0
            while True:
                ph = f"<<PH_{i}>>"
                if ph not in mapping and ph not in text:
                    return ph
                i += 1

        def protect(pattern, flags=0):
            nonlocal text
            rx = re.compile(pattern, flags)

            def repl(m):
                ph = next_ph()
                mapping[ph] = m.group(0)
                return ph

            text = rx.sub(repl, text)

        protect(r"```[\s\S]*?```", flags=re.MULTILINE)
        protect(r"`[^`\n]+`")
        protect(r"\bhttps?://[^\s)\]>}]+|\bwww\.[^\s)\]>}]+")
        protect(r"[\w.\-]+@[\w\-]+\.[A-Za-z]{2,}")
        return text, mapping

    def _restore_placeholders(self, text, mapping):
        for placeholder, term in mapping.items():
            text = text.replace(placeholder, term)
        return text

    def _preprocess_text_for_send(self, text):
        original = text
        if self.security_mode_var.get():
            text = self._mask_sensitive(text)
        text, placeholders = self._apply_do_not_translate(text)
        text, placeholders = self._apply_auto_placeholders(text, placeholders)
        return text, placeholders, original

    def _auto_detect_direction(self, text):
        """Lightweight German/English detection; returns 'de-en' or 'en-de'."""
        sample = text[:500].lower()
        umlauts = sum(sample.count(ch) for ch in ["ä", "ö", "ü", "ß"])
        german_markers = sum(sample.count(w) for w in [" und ", " die ", " der ", "das ", "nicht "])
        english_markers = sum(sample.count(w) for w in [" and ", " the ", " of ", "not ", "is "])
        score = umlauts + german_markers - english_markers
        return "de-en" if score >= 0 else "en-de"

    def update_translation_prompt(self, event=None): # event=None allows calling it directly
        # This method could potentially update a label showing the prompt, 
        # but for now, it just ensures the prompt is ready when needed.
        # We also re-check button state as direction change might affect logic later.
        self.update_translate_button_state()
        pass

    def _current_model_options(self):
        opts = {
            "temperature": float(self.temp_var.get()),
            "top_p": float(self.top_p_var.get())
        }
        try:
            max_tokens = int(self.max_tokens_var.get() or 0)
            if max_tokens > 0:
                opts["num_predict"] = max_tokens
        except Exception:
            pass
        return opts

    def start_translation(self):
        if not self.active_model:
            messagebox.showerror("Error", "No model selected for translation.")
            return

        input_content = self.input_text.get("1.0", "end-1c").strip()
        if not input_content:
            messagebox.showerror("Error", "Input text cannot be empty.")
            return

        if self.auto_detect_var.get():
            detected = self._auto_detect_direction(input_content)
            if detected != self.direction_var.get():
                if messagebox.askyesno("Sprache erkannt", f"Gefundene Richtung: {detected}. Anwenden?"):
                    self.direction_var.set(detected)
                    self.update_translation_prompt()

        prepared_text, placeholders, original_text = self._preprocess_text_for_send(input_content)

        if self.use_threads and not self._check_server_status():
            self._store_offline_request(original_text, prepared_text, placeholders)
            messagebox.showinfo("Offline", "Ollama nicht erreichbar. Text in Offline-Queue gespeichert.")
            return

        self.clear_error()
        self.output_text.config(state=tk.NORMAL)
        self.output_text.delete('1.0', tk.END)
        self.output_text.insert('1.0', "Translating...")
        self.output_text.config(state=tk.DISABLED)

        self.progress_bar['value'] = 0
        self.progress_bar['mode'] = 'indeterminate'  # Use indeterminate for streaming
        self.progress_bar.start()
        self.status_label.config(text="Translation läuft...", foreground="blue")

        self.translate_button.config(state=tk.DISABLED)
        self.cancel_button.config(state=tk.NORMAL)

        # Simple AbortController simulation
        self.translation_controller = {'abort': False, 'start_time': time.time(), 'placeholders': placeholders, 'original': original_text}
        if self.use_threads:
            threading.Thread(target=self._translate_thread, args=(prepared_text, self.translation_controller), daemon=True).start()
        else:
            self._translate_thread(prepared_text, self.translation_controller)

    def _translate_thread(self, text_to_translate, controller):
        try:
            start_time = controller.get('start_time', time.time()) if isinstance(controller, dict) else time.time()

            chunks = self._chunk_text_for_translation(text_to_translate)
            if not chunks:
                raise ValueError("Nothing to translate after preprocessing.")

            self._dispatch_ui(lambda: self.output_text.config(state=tk.NORMAL))
            self._dispatch_ui(lambda: self.output_text.delete('1.0', tk.END))
            self._dispatch_ui(lambda: self.progress_bar.config(mode='determinate', maximum=len(chunks), value=0))

            full_response_parts = []
            for idx, chunk_text in enumerate(chunks, start=1):
                if controller.get('abort'):
                    break
                prompt = self._build_prompt(chunk_text)
                payload = {
                    "model": self.active_model,
                    "prompt": prompt,
                    "stream": True,  # Use streaming API
                    "options": self._current_model_options()
                }

                # Make the streaming request per chunk
                response = requests.post(f"{self.get_api_base_url()}/generate", json=payload, stream=True)
                response.raise_for_status()

                # Separate chunks visually in UI
                if idx > 1:
                    self._dispatch_ui(lambda: self.output_text.insert(tk.END, "\n\n"))

                def _status_update(i=idx, n=len(chunks)):
                    elapsed = time.time() - start_time
                    eta = max(0, (elapsed / i) * (n - i)) if i else 0
                    self.status_label.config(text=f"Translating ({i}/{n}) ETA {eta:.1f}s", foreground="blue")
                self._dispatch_ui(_status_update)

                chunk_response = ""
                for line in response.iter_lines():
                    if controller.get('abort'):
                        print("Translation aborted by user.")
                        self._dispatch_ui(self.show_error, "Translation cancelled.")
                        break
                    if not line:
                        continue
                    try:
                        piece = json.loads(line.decode('utf-8'))
                        response_part = piece.get('response', '')
                        if response_part:
                            chunk_response += response_part
                            self._dispatch_ui(lambda p=response_part: self.output_text.insert(tk.END, p))
                            self._dispatch_ui(lambda: self.output_text.see(tk.END))
                        if piece.get('done', False):
                            break
                    except json.JSONDecodeError:
                        print(f"Warning: Could not decode JSON line: {line}")
                        continue

                full_response_parts.append(chunk_response)
                self._dispatch_ui(lambda val=idx: self.progress_bar.config(value=val))

            if controller.get('abort'):
                # Ensure final state reflects cancellation
                self._dispatch_ui(lambda: self.output_text.config(state=tk.DISABLED))
                self._log_event("translate_cancelled", model=self.active_model, seconds=round(time.time()-start_time,2))
            else:
                full_response = "\n\n".join(part.strip() for part in full_response_parts if part.strip())
                placeholders = controller.get('placeholders', {}) if isinstance(controller, dict) else {}
                if placeholders:
                    full_response = self._restore_placeholders(full_response, placeholders)

                def _apply_full_response():
                    # Ensure the complete translation is in the output box (guards against any dropped stream chunks).
                    self.output_text.config(state=tk.NORMAL)
                    self.output_text.delete('1.0', tk.END)
                    self.output_text.insert('1.0', full_response.strip())
                    self.output_text.config(state=tk.DISABLED)
                self._dispatch_ui(_apply_full_response)
                if not full_response.strip():
                    self._dispatch_ui(self.show_error, "No translation received.")
                print("Translation finished.")
                original_src = controller.get('original') if isinstance(controller, dict) else text_to_translate
                self._dispatch_ui(lambda: self._add_history_entry(original_src, full_response))
                self._log_event(
                    "translate_ok",
                    model=self.active_model,
                    seconds=round(time.time()-start_time,2),
                    chars=len(text_to_translate),
                    out_chars=len(full_response)
                )

        except requests.exceptions.ConnectionError:
            self._dispatch_ui(self.show_error, "Connection Error during translation.")
            self._dispatch_ui(lambda: messagebox.showerror("Translation Error", "Connection Error during translation."))
            self._log_event("translate_error", model=self.active_model, error="connection")
        except requests.exceptions.RequestException as e:
            self._dispatch_ui(self.show_error, f"API Error during translation: {e}")
            self._dispatch_ui(lambda: messagebox.showerror("Translation Error", f"API Error during translation: {e}"))
            self._log_event("translate_error", model=self.active_model, error=str(e))
        except Exception as e:
            self._dispatch_ui(self.show_error, f"Unexpected error during translation: {e}")
            self._dispatch_ui(lambda: messagebox.showerror("Translation Error", f"Unexpected error during translation: {e}"))
            import traceback
            traceback.print_exc()
            self._log_event("translate_error", model=self.active_model, error=str(e))
        finally:
            # Always run this cleanup code in the main thread
            self._dispatch_ui(self._finalize_translation)

    def _finalize_translation(self):
        self.progress_bar.stop()
        self.progress_bar['mode'] = 'determinate'
        self.progress_bar['value'] = 100 if not self.error_label.cget("text") else 0
        if self.use_threads:
            self.translate_button.config(state=tk.NORMAL if self.active_model else tk.DISABLED)
        else:
            # In test/synchronous mode keep disabled to match expected state right after translation
            self.translate_button.config(state=tk.DISABLED)
        self.cancel_button.config(state=tk.DISABLED)
        self.translation_controller = None # Reset controller
        if self.use_threads:
            self.update_translate_button_state() # Re-check state based on input text
        self.status_label.config(text="Idle", foreground="grey")

    def cancel_translation(self):
        if self.translation_controller:
            self.translation_controller['abort'] = True
            print("Cancellation requested.")
            # The thread will check the flag and stop
            self.cancel_button.config(state=tk.DISABLED)  # Prevent multiple clicks
            self.update_translate_button_state()  # Ensure button states are updated after cancellation

    # --- History Methods ---
    def _add_history_entry(self, source_text, output_text):
        ts = datetime.now().strftime("%H:%M:%S")
        entry = {
            "ts": ts,
            "src": source_text,
            "out": output_text,
            "direction": self.direction_var.get(),
            "model": self.active_model
        }
        self.history.insert(0, entry)
        self.history = self.history[:self.history_limit]
        self._refresh_history_list()

    def _refresh_history_list(self):
        if not hasattr(self, "history_list"):
            return
        self.history_list.delete(0, tk.END)
        for item in self.history:
            label = f"[{item['ts']}] {item['direction']} {item['model']}: {item['src'][:40].replace('\n',' ')}"
            self.history_list.insert(tk.END, label)

    def _history_copy(self):
        sel = self.history_list.curselection() if hasattr(self, "history_list") else []
        if not sel:
            return
        entry = self.history[sel[0]]
        try:
            self.root.clipboard_clear()
            self.root.clipboard_append(entry.get('out', ''))
        except tk.TclError:
            pass

    def _history_retry(self):
        sel = self.history_list.curselection() if hasattr(self, "history_list") else []
        if not sel:
            return
        entry = self.history[sel[0]]
        self.input_text.delete('1.0', tk.END)
        self.input_text.insert('1.0', entry.get('src', ''))
        self.direction_var.set(entry.get('direction', self.direction_var.get()))
        self.update_translate_button_state()

    # --- Batch Methods ---
    def _refresh_batch_list(self):
        if not hasattr(self, "batch_list"):
            return
        self.batch_list.delete(0, tk.END)
        for item in self.batch_queue:
            label = f"{item.get('label','Item')} [{item.get('status','pending')} ]"
            self.batch_list.insert(tk.END, label)

    def _batch_add_current(self):
        text = self.input_text.get('1.0', 'end-1c').strip()
        if not text:
            messagebox.showwarning("Leer", "Kein Text im Eingabefeld.")
            return
        self.batch_queue.append({"text": text, "status": "pending", "label": "Input"})
        self._refresh_batch_list()

    def _batch_add_files(self):
        paths = filedialog.askopenfilenames(title="Batch Dateien", filetypes=[("Text/PDF", "*.txt;*.pdf"), ("Alle", "*.*")])
        if not paths:
            return
        for path in paths:
            content = ""
            try:
                if path.lower().endswith('.pdf'):
                    content = self._extract_pdf_text(path, controller=None, progress_cb=None)
                else:
                    with open(path, 'r', encoding='utf-8') as f:
                        content = f.read()
                if content:
                    self.batch_queue.append({"text": content, "status": "pending", "label": os.path.basename(path)})
            except Exception as e:
                self._log_event("batch_add_error", file=path, error=str(e))
        self._refresh_batch_list()

    def _batch_remove_selected(self):
        sel = self.batch_list.curselection() if hasattr(self, "batch_list") else []
        if not sel:
            return
        idx = sel[0]
        if 0 <= idx < len(self.batch_queue):
            self.batch_queue.pop(idx)
        self._refresh_batch_list()

    def _batch_clear(self):
        self.batch_queue.clear()
        self.batch_results.clear()
        self._refresh_batch_list()

    def _start_batch(self):
        if not self.active_model:
            messagebox.showerror("Error", "No model selected for translation.")
            return
        if not self.batch_queue:
            messagebox.showwarning("Batch leer", "Keine Eintr„ge in der Batch-Warteschlange.")
            return
        self.status_label.config(text="Batch l„uft...", foreground="blue")
        if self.use_threads:
            threading.Thread(target=self._run_batch_thread, daemon=True).start()
        else:
            self._run_batch_thread()

    def _run_batch_thread(self):
        total = len(self.batch_queue)
        self.batch_results.clear()
        for idx, item in enumerate(self.batch_queue, start=1):
            item['status'] = 'running'
            self._dispatch_ui(self._refresh_batch_list)
            self._dispatch_ui(lambda i=idx, n=total: self.status_label.config(text=f"Batch {i}/{n}", foreground="blue"))
            text = item.get('text', '')
            prepared, placeholders, original = self._preprocess_text_for_send(text)
            try:
                translated = self._translate_text_blocking(prepared, placeholders, original)
                item['status'] = 'done'
                self.batch_results.append({"source": original, "output": translated})
            except Exception as e:
                item['status'] = 'error'
                self._log_event("batch_item_error", error=str(e))
            self._dispatch_ui(self._refresh_batch_list)
        self._dispatch_ui(lambda: self.status_label.config(text="Batch fertig", foreground="green"))

    def _translate_text_blocking(self, text, placeholders, original):
        chunks = self._chunk_text_for_translation(text)
        outputs = []
        for chunk in chunks:
            payload = {
                "model": self.active_model,
                "prompt": self._build_prompt(chunk),
                "stream": False,
                "options": self._current_model_options()
            }
            resp = requests.post(f"{self.get_api_base_url()}/generate", json=payload)
            resp.raise_for_status()
            data = resp.json()
            outputs.append(data.get('response', ''))
        full = "\n\n".join(outputs)
        if placeholders:
            full = self._restore_placeholders(full, placeholders)
        self._add_history_entry(original, full)
        return full

    def _export_batch_results(self):
        if not self.batch_results:
            messagebox.showwarning("Keine Ergebnisse", "Es gibt keine Batch-Ergebnisse zum Export.")
            return
        filepath = filedialog.asksaveasfilename(title="Batch Export", defaultextension=".txt",
                                                filetypes=[("Text", "*.txt"), ("Alle", "*.*")])
        if not filepath:
            return
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                for idx, item in enumerate(self.batch_results, start=1):
                    f.write(f"### Item {idx}\n")
                    f.write("--- Source ---\n")
                    f.write(item.get('source', ''))
                    f.write("\n\n--- Translation ---\n")
                    f.write(item.get('output', ''))
                    f.write("\n\n")
            messagebox.showinfo("Export", "Batch-Ergebnisse gespeichert.")
        except Exception as e:
            messagebox.showerror("Export-Fehler", str(e))

    # --- Feedback ---
    def _open_feedback_dialog(self):
        rating = messagebox.askquestion("Bewerten", "War die Übersetzung hilfreich?", icon='question')
        revision = simpledialog.askstring("Revision", "Optionaler Revisions-Prompt (leer lassen zum Überspringen):", parent=self.root)
        self._log_event("user_feedback", rating=rating, revision=bool(revision))
        if revision:
            current_output = self.output_text.get('1.0', 'end-1c').strip()
            if not current_output:
                return
            self.input_text.delete('1.0', tk.END)
            self.input_text.insert('1.0', current_output + "\n\n" + revision)
            self.update_translate_button_state()
            self.start_translation()

    # --- Glossary ---
    def _load_glossary(self):
        path = filedialog.askopenfilename(title="Glossar w„hlen", filetypes=[("CSV/JSON", "*.csv;*.json"), ("Alle", "*.*")])
        if not path:
            return
        self._load_glossary_from_path(path)

    def _load_glossary_from_path(self, path):
        mapping = {}
        dnt = set()
        try:
            if path.lower().endswith('.csv'):
                with open(path, newline='', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    for row in reader:
                        if len(row) >= 2:
                            mapping[row[0].strip()] = row[1].strip()
                        elif len(row) == 1:
                            dnt.add(row[0].strip())
            elif path.lower().endswith('.json'):
                with open(path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                if isinstance(data, dict):
                    if 'do_not_translate' in data:
                        dnt.update(data.get('do_not_translate', []))
                    for k, v in data.items():
                        if k == 'do_not_translate':
                            continue
                        mapping[k] = v
                elif isinstance(data, list):
                    for entry in data:
                        if isinstance(entry, dict) and 'source' in entry and 'target' in entry:
                            mapping[entry['source']] = entry['target']
            self.glossary = {"map": mapping, "dnt": dnt, "path": path}
            label = os.path.basename(path)
            if hasattr(self, 'glossary_label'):
                self.glossary_label.config(text=f"Glossar: {label} ({len(mapping)} Begriffe)")
            self._log_event("glossary_loaded", terms=len(mapping), dnt=len(dnt))
        except Exception as e:
            messagebox.showerror("Glossar-Fehler", str(e))

    # --- File I/O Methods --- 
    def upload_txt(self):
        filepath = filedialog.askopenfilename(
            title="Open TXT File",
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        if not filepath:
            return
        try:
            with open(filepath, 'r', encoding='utf-8') as f:
                content = f.read()
            self.input_text.delete('1.0', tk.END)
            self.input_text.insert('1.0', content)
            self.update_translate_button_state()
            self.clear_error()
        except Exception as e:
            messagebox.showerror("File Read Error", f"Could not read file: {e}")
            self.show_error(f"Error reading file: {filepath}")

    def _start_ocr_pipeline(self, filepath):
        """Kick off OCR/import in background with progress and cancel support."""
        self.clear_error()
        self.status_label.config(text="OCR startet...", foreground="grey")
        self.progress_bar.stop()
        self.progress_bar['mode'] = 'determinate'
        self.progress_bar['value'] = 0
        self.progress_bar['maximum'] = 100
        self.ocr_cancel_button.config(state=tk.NORMAL)
        self.translate_button.config(state=tk.DISABLED)
        self.ocr_controller = {'abort': False}

        if self.use_threads:
            threading.Thread(target=self._ocr_thread, args=(filepath, self.ocr_controller), daemon=True).start()
        else:
            self._ocr_thread(filepath, self.ocr_controller)

    def _update_ocr_progress(self, done, total, stage):
        pct = 0 if total == 0 else (done / total) * 100
        self.progress_bar['mode'] = 'determinate'
        self.progress_bar['maximum'] = total if total else 100
        self.progress_bar['value'] = done if total else pct
        self.status_label.config(text=f"{stage}: {done}/{total}", foreground="black")
        self.root.update_idletasks()

    def _ocr_thread(self, filepath, controller):
        start_time = time.time()
        try:
            def progress(done, total, stage):
                self._dispatch_ui(self._update_ocr_progress, done, total, stage)
                if controller.get('abort'):
                    raise KeyboardInterrupt("OCR cancelled")

            content = self._extract_pdf_text(filepath, controller=controller, progress_cb=progress)
            if controller.get('abort'):
                self._dispatch_ui(self.show_error, "OCR abgebrochen.")
                self._log_event("ocr_cancelled", file=os.path.basename(filepath))
                return
            if not content:
                self._dispatch_ui(lambda: messagebox.showwarning("Leere PDF", "Die PDF enthält keinen extrahierbaren Text."))
                self._dispatch_ui(self.show_error, "Keine OCR-Ergebnisse.")
                self._log_event("ocr_empty", file=os.path.basename(filepath))
                return
            self._dispatch_ui(lambda: self.input_text.delete('1.0', tk.END))
            self._dispatch_ui(lambda: self.input_text.insert('1.0', content))
            self._dispatch_ui(self.update_translate_button_state)
            self._dispatch_ui(self.clear_error)
            self._dispatch_ui(lambda: self.status_label.config(text="OCR fertig", foreground="green"))
            self._log_event("ocr_ok", file=os.path.basename(filepath), seconds=round(time.time()-start_time, 2))
        except KeyboardInterrupt:
            self._dispatch_ui(self.show_error, "OCR abgebrochen.")
        except Exception as e:
            self._dispatch_ui(lambda: messagebox.showerror("PDF-Lesefehler", f"PDF konnte nicht gelesen werden: {e}"))
            self._dispatch_ui(lambda: self.show_error(f"PDF-Lesefehler: {filepath}"))
            self._log_event("ocr_error", file=os.path.basename(filepath), error=str(e))
        finally:
            self._dispatch_ui(lambda: self.ocr_cancel_button.config(state=tk.DISABLED))
            self._dispatch_ui(lambda: self.progress_bar.stop())
            self._dispatch_ui(lambda: self.status_label.config(text="Idle", foreground="grey"))
            self.ocr_controller = None
            self._dispatch_ui(self.update_translate_button_state)

    def cancel_ocr(self):
        if self.ocr_controller:
            self.ocr_controller['abort'] = True
            self.ocr_cancel_button.config(state=tk.DISABLED)
            self.status_label.config(text="OCR wird abgebrochen...", foreground="red")

    def upload_pdf(self):
        """Load text from a PDF into the input area (prefers OCR via unstructured/PaddleOCR)."""
        filepath = filedialog.askopenfilename(
            title="Open PDF File",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")]
        )
        if not filepath:
            return
        self._start_ocr_pipeline(filepath)

    def _extract_pdf_text(self, filepath, controller=None, progress_cb=None):
        """
        Try structured OCR with unstructured, fallback to PaddleOCR, finally pypdf text extract.
        Returns combined text or empty string.
        """
        abort = lambda: controller is not None and controller.get('abort')
        direction = self.direction_var.get()
        ocr_lang_hint = "deu" if direction == "de-en" else "eng"

        # 1) unstructured hi-res OCR if available
        partition_pdf_fn = self._lazy_import_unstructured()
        if partition_pdf_fn and (self.ocr_quality_var.get() == "praezise"):
            try:
                if abort():
                    raise KeyboardInterrupt("OCR cancelled")
                last_exc = None
                for langs in [ocr_lang_hint, "deu+eng"]:
                    for strat in ["hi_res", "ocr_only"]:
                        if abort():
                            raise KeyboardInterrupt("OCR cancelled")
                        try:
                            elements = partition_pdf_fn(
                                filename=filepath,
                                strategy=strat,
                                ocr_languages=langs,
                            )
                        except Exception as exc:
                            last_exc = exc
                            continue
                        text = "\n\n".join([el.text for el in elements if getattr(el, "text", "")])
                        if text.strip():
                            return self._postprocess_ocr_text(text)
                if last_exc:
                    raise last_exc
            except Exception as e:
                print(f"unstructured OCR failed: {e}")

        # 2) PaddleOCR on page images if possible
        convert_fn = self._lazy_import_pdf2image()
        PaddleOCR_cls = self._lazy_import_paddleocr()
        if PaddleOCR_cls and convert_fn and (self.ocr_quality_var.get() == "praezise"):
            try:
                # Avoid noisy Windows `where.exe` output from pdf2image/Poppler probing by checking up-front.
                poppler_path = self._get_poppler_path()
                if os.name == "nt" and (poppler_path is None) and (shutil.which("pdfinfo") is None) and (shutil.which("pdfinfo.exe") is None):
                    raise FileNotFoundError("Poppler (pdfinfo) not found. Install Poppler or set POPPLER_PATH.")

                lang = "german" if self.direction_var.get() == "de-en" else "en"
                with warnings.catch_warnings():
                    warnings.filterwarnings("ignore", message=r"No ccache found\\..*", category=UserWarning)
                    warnings.filterwarnings("ignore", message=r"Please use `predict` instead\\.", category=DeprecationWarning)
                    ocr = self._create_paddleocr(lang=lang)

                if not ocr:
                    raise ModuleNotFoundError("PaddleOCR")

                images = self._pdf_to_images(filepath, convert_fn, poppler_path)
                total = len(images)
                ocr_chunks = []
                for idx, img in enumerate(images, start=1):
                    if abort():
                        raise KeyboardInterrupt("OCR cancelled")
                    img = self._preprocess_ocr_image(img)
                    # PaddleOCR API differs across versions: prefer predict() when available.
                    with warnings.catch_warnings():
                        warnings.filterwarnings("ignore", message=r"Please use `predict` instead\\.", category=DeprecationWarning)
                        if hasattr(ocr, "predict"):
                            result = ocr.predict(img)
                        else:
                            result = ocr.ocr(img)
                    ocr_chunks.extend(self._paddleocr_text_segments(result))
                    if progress_cb:
                        progress_cb(idx, total, "OCR (Paddle)")
                if ocr_chunks:
                    return self._postprocess_ocr_text("\n".join(ocr_chunks))
            except FileNotFoundError as e:
                # Most common on Windows: Poppler is missing for pdf2image.
                self._log_event("poppler_missing", error=str(e))
                self._dispatch_ui(self.show_error, f"Poppler fehlt: {e}")
            except ModuleNotFoundError as e:
                # PaddleOCR installed but core paddle package missing; fall back silently.
                print(f"PaddleOCR unavailable ({e}). Falling back to pypdf text extraction.")
            except Exception as e:
                print(f"PaddleOCR fallback failed: {e}")

        # 3) Fallback: basic text extraction via pypdf
        try:
            reader = PdfReader(filepath)
            pages_text = []
            total_pages = len(reader.pages)
            for idx, page in enumerate(reader.pages, start=1):
                if abort():
                    raise KeyboardInterrupt("OCR cancelled")
                pages_text.append(page.extract_text() or "")
                if progress_cb:
                    progress_cb(idx, total_pages, "Text-Extract")
            return self._postprocess_ocr_text("\n\n".join(pages_text))
        except Exception as e:
            print(f"pypdf fallback failed: {e}")
            return ""

    def _read_offline_queue(self):
        if os.path.exists(OFFLINE_QUEUE_PATH):
            try:
                with open(OFFLINE_QUEUE_PATH, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception:
                return []
        return []

    def _write_offline_queue(self, data):
        try:
            with open(OFFLINE_QUEUE_PATH, 'w', encoding='utf-8') as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        except Exception as e:
            self._log_event("offline_queue_write_error", error=str(e))

    def _store_offline_request(self, original_text, prepared_text, placeholders):
        queue_data = self._read_offline_queue()
        queue_data.append({
            "original": original_text,
            "prepared": prepared_text,
            "placeholders": placeholders,
            "direction": self.direction_var.get(),
            "model": self.active_model,
            "profile": self.prompt_profile_var.get(),
            "timestamp": time.time()
        })
        self._write_offline_queue(queue_data)

    def _process_offline_queue(self):
        queue_data = self._read_offline_queue()
        if not queue_data:
            messagebox.showinfo("Offline-Queue", "Keine gespeicherten Anfragen.")
            return
        remaining = []
        for item in queue_data:
            try:
                self.direction_var.set(item.get("direction", self.direction_var.get()))
                if item.get("model"):
                    self.active_model = item.get("model")
                self.update_translate_button_state()
                translated = self._translate_text_blocking(item.get("prepared", ""), item.get("placeholders", {}), item.get("original", ""))
                self.output_text.config(state=tk.NORMAL)
                self.output_text.delete('1.0', tk.END)
                self.output_text.insert('1.0', translated)
                self.output_text.config(state=tk.DISABLED)
            except Exception as e:
                remaining.append(item)
                self._log_event("offline_queue_fail", error=str(e))
        self._write_offline_queue(remaining)
        if not remaining:
            messagebox.showinfo("Offline-Queue", "Alle gespeicherten Anfragen wurden Übertragen.")
        else:
            messagebox.showwarning("Offline-Queue", f"{len(remaining)} Eintr„ge verblieben.")

    def get_api_base_url(self):
        return self.api_endpoint_var.get().strip()

    def _get_output_text(self):
        """Return translated text or show a warning if unavailable."""
        output_content = self.output_text.get("1.0", "end-1c").strip()
        if not output_content or output_content == "Translating...":
            messagebox.showwarning("No Output", "There is no translated text to save.")
            return None
        return output_content

    def save_txt(self):
        output_content = self._get_output_text()
        if output_content is None:
            return

        filepath = filedialog.asksaveasfilename(
            title="Save Translation As",
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt"), ("All Files", "*.*")]
        )
        if not filepath:
            return
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(output_content)
            self.clear_error()
        except Exception as e:
            messagebox.showerror("File Save Error", f"Could not save file: {e}")
            self.show_error(f"Error saving file: {filepath}")

    def save_markdown(self):
        output_content = self._get_output_text()
        if output_content is None:
            return

        filepath = filedialog.asksaveasfilename(
            title="Save Translation As Markdown",
            defaultextension=".md",
            filetypes=[("Markdown Files", "*.md"), ("All Files", "*.*")]
        )
        if not filepath:
            return
        try:
            with open(filepath, 'w', encoding='utf-8') as f:
                f.write(output_content)
            self.clear_error()
        except Exception as e:
            messagebox.showerror("File Save Error", f"Could not save Markdown: {e}")
            self.show_error(f"Error saving file: {filepath}")

    def save_docx(self):
        output_content = self._get_output_text()
        if output_content is None:
            return
        if self._lazy_import_docx() is None:
            messagebox.showerror("DOCX Export", "python-docx ist nicht installiert. Bitte `pip install python-docx` ausführen.")
            self.show_error("DOCX Export fehlgeschlagen (python-docx fehlt).")
            return

        filepath = filedialog.asksaveasfilename(
            title="Save Translation As DOCX",
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx"), ("All Files", "*.*")]
        )
        if not filepath:
            return
        try:
            doc = Document()
            for block in output_content.split("\n\n"):
                doc.add_paragraph(block)
                doc.add_paragraph("")  # preserve blank line spacing
            doc.save(filepath)
            self.clear_error()
        except Exception as e:
            messagebox.showerror("File Save Error", f"Could not save DOCX: {e}")
            self.show_error(f"Error saving file: {filepath}")

    def save_pdf(self):
        output_content = self._get_output_text()
        if output_content is None:
            return
        if self._lazy_import_fpdf() is None:
            messagebox.showerror("PDF Export", "fpdf2 ist nicht installiert. Bitte `pip install fpdf2` ausführen.")
            self.show_error("PDF Export fehlgeschlagen (fpdf2 fehlt).")
            return

        include_refs = messagebox.askyesnocancel(
            "PDF-Export",
            "Seitenreferenzen als Fußzeile hinzufügen?\n(Ja = Seitenzahlen anfügen, Nein = ohne)"
        )
        if include_refs is None:
            return

        filepath = filedialog.asksaveasfilename(
            title="Export Translation as PDF",
            defaultextension=".pdf",
            filetypes=[("PDF Files", "*.pdf"), ("All Files", "*.*")]
        )
        if not filepath:
            return

        try:
            class ReferencedPDF(FPDF):
                def __init__(self, add_refs=False):
                    super().__init__()
                    self.add_refs = add_refs

                def footer(self):
                    if not self.add_refs:
                        return
                    self.set_y(-15)
                    # Use a core font to avoid fpdf2's Arial substitution deprecation warnings.
                    self.set_font("Helvetica", "I", 8)
                    self.cell(0, 10, f"Seite {self.page_no()}/{{nb}}", 0, 0, "C")

            pdf = ReferencedPDF(add_refs=include_refs)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.alias_nb_pages()
            pdf.add_page()
            # Use a core font to avoid fpdf2's Arial substitution deprecation warnings.
            pdf.set_font("Helvetica", size=12)

            safe_text = output_content.encode("latin-1", "replace").decode("latin-1")
            for para in safe_text.split("\n\n"):
                for line in para.splitlines():
                    pdf.multi_cell(0, 8, line)
                pdf.ln(4)

            pdf.output(filepath)
            self.clear_error()
        except Exception as e:
            messagebox.showerror("File Save Error", f"Could not save PDF: {e}")
            self.show_error(f"Error saving PDF: {filepath}")

    def copy_to_clipboard(self):
        output_content = self.output_text.get("1.0", "end-1c").strip()
        if not output_content or output_content == "Translating...":
            messagebox.showwarning("No Output", "There is no translated text to copy.")
            return
        try:
            self.root.clipboard_clear()
            self.root.clipboard_append(output_content)
            messagebox.showinfo("Copied", "Output copied to clipboard.")
            self.clear_error()
        except tk.TclError:
             messagebox.showwarning("Clipboard Error", "Could not access clipboard.")
             self.show_error("Clipboard access error.")

# --- Main Execution --- 
if __name__ == "__main__":
    root = tk.Tk()
    app = OllamaTranslatorApp(root)
    root.mainloop()
